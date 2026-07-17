"""知识库文档服务：上传、解析、分块、向量化、存储"""
import asyncio
import os
from io import BytesIO

import chromadb
import httpx
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config.knowledge_base import get_kb_settings
from config.database import SessionLocal
from models.knowledge_base import KnowledgeDocument

kb_settings = get_kb_settings()

_chroma_path = os.path.abspath(kb_settings.CHROMA_PERSIST_DIR)
os.makedirs(_chroma_path, exist_ok=True)
_chroma_client = chromadb.PersistentClient(path=_chroma_path)


class DashScopeEmbedding:
    """阿里云百炼 Embedding 封装（httpx 直调，避免 LangChain OpenAIEmbeddings 兼容问题）"""

    def __init__(self):
        self.api_key = kb_settings.EMBEDDING_API_KEY
        self.base_url = kb_settings.EMBEDDING_BASE_URL.rstrip("/")
        self.model = kb_settings.EMBEDDING_MODEL

    def embed_query(self, text: str) -> list[float]:
        return self.embed_documents([text])[0]

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        resp = httpx.post(
            f"{self.base_url}/embeddings",
            json={"model": self.model, "input": texts},
            headers={"Authorization": f"Bearer {self.api_key}"},
            timeout=60.0,
        )
        data = resp.json()
        if resp.status_code != 200:
            raise RuntimeError(f"Embedding API error: {data}")
        items = sorted(data.get("data", []), key=lambda x: x.get("index", 0))
        return [item["embedding"] for item in items]


embeddings = DashScopeEmbedding()


def _get_collection(user_id: int):
    return _chroma_client.get_or_create_collection(name=f"user_{user_id}_kb")


def _parse_file(file_data: bytes, filename: str) -> str | None:
    """解析文件内容，返回纯文本"""
    ext = os.path.splitext(filename)[1].lower()
    try:
        if ext == ".txt" or ext == ".md":
            return file_data.decode("utf-8", errors="replace")
        elif ext == ".pdf":
            import fitz
            doc = fitz.open(stream=file_data, filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        elif ext == ".docx":
            from docx import Document
            doc = Document(BytesIO(file_data))
            return "\n".join(p.text for p in doc.paragraphs)
    except Exception:
        return None
    return None


def upload_and_process(user_id: int, file_data: bytes, filename: str) -> int | None:
    """上传并处理文档，返回 document_id"""
    db = SessionLocal()
    try:
        count = db.query(KnowledgeDocument).filter(
            KnowledgeDocument.user_id == user_id
        ).count()
        if count >= kb_settings.MAX_DOCUMENTS_PER_USER:
            return None
    finally:
        db.close()

    text = _parse_file(file_data, filename)
    if not text:
        return None

    if len(text) > kb_settings.MAX_TEXT_LENGTH:
        text = text[:kb_settings.MAX_TEXT_LENGTH]

    file_type = os.path.splitext(filename)[1].lower().lstrip(".")
    file_size = len(file_data)

    db = SessionLocal()
    try:
        doc = KnowledgeDocument(
            user_id=user_id,
            title=filename,
            file_type=file_type,
            file_size=file_size,
            chunk_count=0,
            status="processing",
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)
        doc_id = doc.id
    finally:
        db.close()

    asyncio.create_task(_process_document(doc_id, text, user_id))
    return doc_id


async def _process_document(doc_id: int, text: str, user_id: int):
    """后台处理：分块、向量化、存储"""
    db = SessionLocal()
    try:
        doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
        if not doc:
            return

        try:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=kb_settings.CHUNK_SIZE,
                chunk_overlap=kb_settings.CHUNK_OVERLAP,
                separators=["\n\n", "\n", "。", ".", " ", ""],
            )
            chunks = splitter.split_text(text)

            if not chunks:
                doc.status = "completed"
                doc.chunk_count = 0
                db.commit()
                return

            collection = _get_collection(user_id)
            chunk_metas = []
            chunk_ids = []
            for i, chunk in enumerate(chunks):
                chunk_metas.append({"document_id": str(doc_id), "title": doc.title})
                chunk_ids.append(f"doc_{doc_id}_chunk_{i}")

            # 批量向量化（每次最多 20 条，避免超时）
            all_embeddings = []
            batch_size = 20
            for i in range(0, len(chunks), batch_size):
                batch = chunks[i:i + batch_size]
                all_embeddings.extend(embeddings.embed_documents(batch))

            collection.add(
                embeddings=all_embeddings,
                documents=chunks,
                metadatas=chunk_metas,
                ids=chunk_ids,
            )

            doc.status = "completed"
            doc.chunk_count = len(chunks)
            db.commit()
        except Exception as e:
            doc.status = "failed"
            db.commit()
            print(f"[KB] Process error: {e}")
    finally:
        db.close()


def delete_document(user_id: int, doc_id: int) -> bool:
    """删除文档（DB + ChromaDB）"""
    db = SessionLocal()
    try:
        doc = db.query(KnowledgeDocument).filter(
            KnowledgeDocument.id == doc_id,
            KnowledgeDocument.user_id == user_id,
        ).first()
        if not doc:
            return False

        try:
            collection = _get_collection(user_id)
            chunk_ids = [f"doc_{doc_id}_chunk_{i}" for i in range(doc.chunk_count + 10)]
            collection.delete(ids=chunk_ids)
        except Exception:
            pass

        db.delete(doc)
        db.commit()
        return True
    finally:
        db.close()


def list_documents(user_id: int) -> list[dict]:
    """列出用户的所有文档"""
    db = SessionLocal()
    try:
        docs = db.query(KnowledgeDocument).filter(
            KnowledgeDocument.user_id == user_id
        ).order_by(KnowledgeDocument.create_at.desc()).all()
        return [
            {
                "id": d.id,
                "title": d.title,
                "file_type": d.file_type,
                "file_size": d.file_size,
                "status": d.status,
                "created_at": d.create_at.isoformat() if d.create_at else None,
            }
            for d in docs
        ]
    finally:
        db.close()


def get_document_count(user_id: int) -> int:
    """获取用户文档数量"""
    db = SessionLocal()
    try:
        return db.query(KnowledgeDocument).filter(
            KnowledgeDocument.user_id == user_id,
            KnowledgeDocument.status == "completed",
        ).count()
    finally:
        db.close()
